import os, io, uuid, re, json, datetime as dt
import logging
from typing import List, Optional
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import Response, PlainTextResponse
import pandas as pd

# Optional deps
try:
    import docx  # python-docx
except Exception:
    docx = None
try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None

EXPORT_DIR = os.environ.get("EXPORT_DIR", "/data/exports")
os.makedirs(EXPORT_DIR, exist_ok=True)

app = FastAPI(title="AI-ДОВЕРИЕ API", version="1.0")



# --- CORS: configurable via env (supports list and regex) ---
logger = logging.getLogger("uvicorn.error")
_cors_allow_origins_env = os.environ.get("CORS_ALLOW_ORIGINS", "")
_cors_allow_origin_regex_env = os.environ.get("CORS_ALLOW_ORIGIN_REGEX", "")
_allow_origins = [o.strip() for o in _cors_allow_origins_env.split(",") if o.strip()]
_allow_origin_regex = None
if _cors_allow_origin_regex_env.strip():
    try:
        _allow_origin_regex = re.compile(_cors_allow_origin_regex_env.strip())
    except re.error as e:
        logger.warning(f"Invalid CORS regex '{_cors_allow_origin_regex_env}': {e}")

_allow_credentials = os.environ.get("CORS_ALLOW_CREDENTIALS", "false").lower() == "true"
if "*" in _allow_origins and _allow_credentials:
    _allow_credentials = False

app.add_middleware(
    CORSMiddleware,
    allow_origins=_allow_origins if _allow_origins else [],
    allow_origin_regex=_allow_origin_regex.pattern if _allow_origin_regex else None,
    allow_credentials=_allow_credentials,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)

CATEGORIES = [
 "Благоустройство","Окружающая среда","Доступность цифровых услуг","Дороги","Образование",
 "Культура","Здравоохранение","Транспортное обслуживание","ЖКХ","Адаптация участников СВО","Политическое доверие"
]

MUNICIPALITIES = [
    {"id":1,"name":"Раменский"},
    {"id":2,"name":"Жуковский"},
    {"id":3,"name":"Бронницы"},
    {"id":4,"name":"Люберцы"},
]

KEYWORDS = {
    "Благоустройство": ["дворы","освещение","урны","лавочки","парк","сквер","уборка","детская площадка","озеленение","благоустройство"],
    "Окружающая среда": ["экология","свалка","запах","дым","выбросы","река","водоём","шум","окружающая среда","природа"],
    "Доступность цифровых услуг": ["госуслуги","интернет","цифров","сайт","онлайн","мфц запись","портал"],
    "Дороги": ["дорога","ямы","ремонт дороги","асфальт","яма","бордюр","разметка","снег","уборка снега","тротуар"],
    "Образование": ["школа","детсад","садик","учитель","образование","лицей","гимназия"],
    "Культура": ["культура","дом культуры","библиотека","музей","концерт"],
    "Здравоохранение": ["поликлиника","больница","врач","медицина","здравоохранение","скорая"],
    "Транспортное обслуживание": ["автобус","маршрут","транспорт","расписание","остановка","электричка","метро"],
    "ЖКХ": ["жкх","квартира","подъезд","управляющая компания","счетчик","отопление","вода","горячая вода","холодная вода","электричество","лифт"],
    "Адаптация участников СВО": ["СВО","ветеран","реабилитация","поддержка","пособие"],
    "Политическое доверие": ["мэр","глава","администрация","власть","политика","доверие"]
}

def guess_category(text:str)->str:
    t = text.lower()
    best = None; score = 0
    for cat, kws in KEYWORDS.items():
        s = sum(1 for w in kws if w in t)
        if s>score: score=s; best=cat
    return best or "ЖКХ"

DATE_RE = re.compile(r"(20\d{2}[-./]\d{1,2}[-./]\d{1,2}|\d{1,2}[-./]\d{1,2}[-./]20\d{2})")
ADDR_RE = re.compile(r"(ул\.\s*[А-ЯЁа-яёA-Za-z0-9\- ]+|проспект\s+[А-ЯЁа-яёA-Za-z\- ]+|дом\s*\d+[А-Яа-яA-Za-z]?)")


DOBRODEL_STATUS_ALLOW = {
    "в работе исполнителя",
    "на исполнении",
    "на уточнении модератора",
}

def parse_dobrodel_excel(upload: UploadFile):
    """Return list of normalized rows from a Добродел выгрузка, or None if not applicable.
    Each row: source,date,address,text,category,lat,lng,municipality_id(None for now)"""
    try:
        xls = pd.ExcelFile(upload.file)
        # choose sheet that has typical columns
        target_sheet = None
        for s in xls.sheet_names:
            cols = [str(c).strip().lower() for c in xls.parse(s, nrows=1).columns]
            if ("омсу" in cols) and (("статус" in cols) or any("статус" in c for c in cols)):
                target_sheet = s
                break
        if not target_sheet:
            return None
        df = xls.parse(target_sheet)
        # normalize columns
        cols_map = {str(c).strip().lower(): c for c in df.columns}
        def pick(*variants):
            for v in variants:
                v_low = v.lower()
                if v_low in cols_map: return cols_map[v_low]
                # fuzzy contains match
                for k in cols_map:
                    if v_low in k: return cols_map[k]
            return None

        col_omcu = pick("ОМСУ")
        col_status = pick("Статус")
        col_source = pick("Источник")
        col_date = pick("Дата обращения","Дата (первого взятия в работу)","Дата")
        col_address = pick("Адрес")
        col_fact = pick("Факт")
        col_descr = pick("Описание")

        if not (col_omcu and col_status and (col_fact or col_descr) and col_address):
            # not a recognizable layout
            return None

        df2 = df.copy()
        # filter OМСУ содержит Люберцы
        df2 = df2[df2[col_omcu].astype(str).str.contains("Люберц", case=False, na=False)]
        # filter statuses
        df2 = df2[df2[col_status].astype(str).str.strip().str.lower().isin(DOBRODEL_STATUS_ALLOW)]
        if df2.empty:
            return []

        # build rows
        rows = []
        for _, r in df2.iterrows():
            src = str(r.get(col_source, "Добродел") or "Добродел")
            date_val = r.get(col_date, None)
            # normalize date to yyyy-mm-dd
            date_str = None
            if pd.notna(date_val):
                try:
                    date_parsed = pd.to_datetime(date_val)
                    date_str = str(date_parsed.date())
                except Exception:
                    date_str = str(date_val)[:10]
            address = str(r.get(col_address, "") or "").strip() or None
            text_parts = []
            if col_fact: 
                v = r.get(col_fact, None)
                if pd.notna(v): text_parts.append(str(v))
            if col_descr:
                v = r.get(col_descr, None)
                if pd.notna(v): text_parts.append(str(v))
            text = "\n".join(text_parts).strip()
            cat = guess_category(text or "")
            # naive geotag placeholder: none (could be enhanced later)
            lat_val, lng_val = detect_coords_from_row(r)
            rows.append({
                "source": src,
                "date": date_str,
                "address": address,
                "text": text,
                "category": cat,
                "lat": lat_val,
                "lng": lng_val,
                "municipality_id": None
            })
        return rows
    except Exception as e:
        # If anything goes wrong, fall back to generic extraction
        return None
def extract_text_from_file(up: UploadFile) -> str:
    name = up.filename or "file"
    if name.lower().endswith((".xlsx",".xls",".csv")):
        try:
            if name.lower().endswith(".csv"):
                df = pd.read_csv(up.file)
            else:
                df = pd.read_excel(up.file)
            cols = [c.lower() for c in df.columns]
            text_cols = [i for i,c in enumerate(cols) if any(x in c for x in ["текст","сообщ","опис","обращ","post","message"])]
            if text_cols:
                return "\n".join(str(x) for x in df.iloc[:, text_cols[0]].astype(str).tolist())
            else:
                return df.to_csv(index=False)
        except Exception as e:
            return f"Не удалось прочитать таблицу: {e}"
    if name.lower().endswith((".doc",".docx")) and docx:
        try:
            d = docx.Document(up.file)
            return "\n".join(p.text for p in d.paragraphs)
        except Exception as e:
            return f"Не удалось прочитать DOCX: {e}"
    if name.lower().endswith(".pdf") and PdfReader:
        try:
            r = PdfReader(up.file)
            texts = []
            for p in r.pages:
                texts.append(p.extract_text() or "")
            return "\n".join(texts)
        except Exception as e:
            return f"Не удалось прочитать PDF: {e}"
    # fallback
    b = up.file.read()
    try:
        return b.decode("utf-8")
    except Exception:
        return b.decode("latin-1","ignore")

def extract_fields(text:str, source:str):
    date_match = DATE_RE.search(text)
    date = None
    if date_match:
        raw = date_match.group(0).replace('/','-').replace('.','-')
        parts = raw.split('-')
        if len(parts[0])==4:
            date = raw
        else:
            # dd-mm-yyyy
            d,m,y = parts
            date = f"{y}-{m.zfill(2)}-{d.zfill(2)}"
    addr_match = ADDR_RE.search(text)
    address = addr_match.group(0) if addr_match else None
    category = guess_category(text)
    lat, lng = detect_coords_from_text((text or "") + " " + (address or ""))
    return {
        "source": source,
        "date": date,
        "address": address,
        "text": text.strip()[:5000],
        "category": category,
        "lat": lat, "lng": lng,
    }

# In-memory "DB"

# --- Geotag detection helpers ---
COORD_DD_RE = re.compile(r'(?P<lat>[+-]?\d{1,2}(?:[.,]\d+))\s*[,; ]\s*(?P<lng>[+-]?\d{1,3}(?:[.,]\d+))')
DMS_RE = re.compile(r'(?:(?P<lat_deg>\d{1,2})[°\s]\s*(?P<lat_min>\d{1,2})(?:[\'’′]\s*(?P<lat_sec>\d{1,2}(?:[.,]\d+))?)?\s*(?P<lat_hem>[NSСЮСеверЮж])\s*[,; ]\s*)?(?P<lng_deg>\d{1,3})[°\s]\s*(?P<lng_min>\d{1,2})(?:[\'’′]\s*(?P<lng_sec>\d{1,2}(?:[.,]\d+))?)?\s*(?P<lng_hem>[EWЗВВостЗап])', re.IGNORECASE)

def _dms_to_dd(deg, minutes, seconds, hemisphere):
    deg = float(str(deg).replace(',', '.'))
    minutes = float(str(minutes).replace(',', '.')) if minutes is not None else 0.0
    seconds = float(str(seconds).replace(',', '.')) if seconds is not None else 0.0
    dd = deg + minutes/60.0 + seconds/3600.0
    if hemisphere and str(hemisphere).upper() in ['S','W','З','Ю']:
        dd = -dd
    return dd

def detect_coords_from_text(text: str):
    if not text:
        return (None, None)
    t = str(text)
    m = COORD_DD_RE.search(t)
    if m:
        try:
            lat = float(m.group('lat').replace(',', '.'))
            lng = float(m.group('lng').replace(',', '.'))
            if -90 <= lat <= 90 and -180 <= lng <= 180:
                return (lat, lng)
        except Exception:
            pass
    m = DMS_RE.search(t)
    if m:
        try:
            lat = None
            if m.group('lat_deg') is not None:
                lat = _dms_to_dd(m.group('lat_deg'), m.group('lat_min'), m.group('lat_sec'), m.group('lat_hem') or 'N')
            lng = _dms_to_dd(m.group('lng_deg'), m.group('lng_min'), m.group('lng_sec'), m.group('lng_hem') or 'E')
            return (lat, lng)
        except Exception:
            pass
    return (None, None)

def detect_coords_from_row(row: dict):
    # Try explicit lat/lng columns
    for key in list(row.keys()):
        k = (str(key) or '').strip().lower()
        if k in ('lat', 'latitude', 'широта', 'y'):
            val = row.get(key)
            if val not in (None, ''):
                try:
                    lat = float(str(val).replace(',', '.'))
                except Exception:
                    continue
                for kk in list(row.keys()):
                    k2 = (str(kk) or '').strip().lower()
                    if k2 in ('lng', 'lon', 'long', 'longitude', 'долгота', 'x'):
                        v2 = row.get(kk)
                        if v2 not in (None, ''):
                            try:
                                lng = float(str(v2).replace(',', '.'))
                                return (lat, lng)
                            except Exception:
                                pass
    # Fallback: detect inside text/address
    joined = " ".join(str(row.get(k) or '') for k in row.keys())
    return detect_coords_from_text(joined)
DB = {
    "rows": [],
    "plans": [],
}

@app.get("/api/appeals/municipalities")
def municipalities():
    return {"items": MUNICIPALITIES}

from fastapi import Request

@app.post("/api/appeals/upload")

@app.post("/api/appeals/upload")
async def upload_appeals(request: Request, files: List[UploadFile] = File(...), municipality_id: Optional[int] = Form(None)):
    if not files:
        raise HTTPException(400, "Файлы не переданы")
    rows = []
    for f in files:
        # Try special Добродел parser for Excel
        parsed = None
        name = (f.filename or '').lower()
        if name.endswith(('.xlsx','.xls')):
            parsed = parse_dobrodel_excel(f)
            f.file.seek(0)
        if isinstance(parsed, list):
            for item in parsed:
                item['municipality_id'] = municipality_id
            rows.extend(parsed)
        else:
            # Fallback: treat whole file as one text blob
            text = extract_text_from_file(f)
            fields = extract_fields(text, f.filename)
            fields["municipality_id"] = municipality_id
            rows.append(fields)
    DB["rows"].extend(rows)

    df = pd.DataFrame(rows, columns=["source","date","address","text","category","lat","lng","municipality_id"])
    export_id = str(uuid.uuid4())
    xlsx_path = os.path.join(EXPORT_DIR, f"{export_id}.xlsx")
    df.to_excel(xlsx_path, index=False)

    origin = str(request.base_url).rstrip('/')
    return {"items": rows, "export_url": f"{origin}/api/appeals/export/{export_id}.xlsx"}

@app.get("/api/appeals/export/{file_name}")
def export_file(file_name:str):
    path = os.path.join(EXPORT_DIR, file_name)
    if not os.path.exists(path):
        raise HTTPException(404, "Файл не найден")
    return FileResponse(path, filename=file_name, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# === Analytics helpers and endpoint ===
RU_STOP = set("и,в,во,не,что,он,на,я,с,со,как,а,то,все,она,так,его,но,да,ты,к,у,же,вы,за,бы,по,ее,мне,есть,тут,они,мы,тебя,ничего,чтобы,когда,где,даже,или,если,без,из,под,при,для,над,про,после,между,это,этот,эта,эти,того,той,тем,теми,тех,та,тут,там,быть,будет,был,были,будут,же,ли,до,от,ну".split(","))

POS_WORDS = set("хорошо,исправили,починили,спасибо,благодарим,улучшили,решено,устранено".split(","))
NEG_WORDS = set("плохо,ужас,проблема,жалоба,не работает,сломано,грязь,мусор,яма,вонь,шум,некачественно,затопило,отсутствует,протечка,нет,нарушение".split(","))

def sentiment_score(text:str)->float:
    t = (text or '').lower()
    score = 0
    for w in POS_WORDS:
        if w in t: score += 1
    for w in NEG_WORDS:
        if w in t: score -= 1
    return 1.0 if score>0 else (-1.0 if score<0 else 0.0)

def top_tokens(texts:List[str], topn:int=5)->List[str]:
    from collections import Counter
    cnt = Counter()
    for t in texts:
        words = re.findall(r"[А-ЯЁа-яёA-Za-z0-9\-]{3,}", (t or "").lower())
        for w in words:
            if w in RU_STOP: continue
            cnt[w]+=1
    return [w for w,_ in cnt.most_common(topn)]

@app.get("/api/appeals/analytics")
def analytics(municipality_id: Optional[int] = None):
    df = pd.DataFrame(DB["rows"] or [], columns=["source","date","address","text","category","lat","lng","municipality_id"])
    if municipality_id:
        df = df[df["municipality_id"]==municipality_id]
    if df.empty:
        return {"by_category": [], "by_date": [], "per_category": []}

    vc = df["category"].fillna("—").astype(str).value_counts()
    by_cat_df = vc.rename_axis("name").reset_index(name="value")
    by_cat = [{"name": str(r["name"]), "value": int(r["value"])} for _, r in by_cat_df.iterrows()]

    df2 = df.copy()
    df2["date"] = df2["date"].fillna("").astype(str).str.slice(0,10).replace({"": "—"})
    vc2 = df2["date"].value_counts()
    by_date_df = vc2.rename_axis("date").reset_index(name="count")
    by_date = [{"date": str(r["date"]), "count": int(r["count"])} for _, r in by_date_df.iterrows()]
    by_date = sorted(by_date, key=lambda x: x["date"])

    per_category = []
    for cat, g in df.groupby("category"):
        texts = g["text"].astype(str).tolist()
        addr_counts = g["address"].dropna().astype(str).value_counts().rename_axis("address").reset_index(name="count")
        hotspots = [{"address": a, "count": int(c)} for a,c in addr_counts.head(5).itertuples(index=False)]
        topics = top_tokens(texts, topn=7)
        s = sum(sentiment_score(t) for t in texts)
        sentiment = round(s / max(1,len(texts)), 3)
        per_category.append({
            "category": str(cat),
            "count": int(len(g)),
            "unique_texts": int(len(g)),
            "hotspots": hotspots,
            "topics": topics,
            "sentiment": sentiment
        })

    return {"by_category": by_cat, "by_date": by_date, "per_category": per_category}
def make_plan_text(category:str, municipality_name:str)->str:
    today = dt.date.today()
    deadline = today + dt.timedelta(days=30)
    steps = [
      ("Диагностика проблематики","Собрать первичные данные, верифицировать адресные точки, составить карту очагов."),
      ("Быстрые победы (до 2 недель)","Отработать 2–3 адреса с высокой видимостью; подготовить фото «до/после»."),
      ("Системные меры","Запланировать закупки/МКУ/подрядчики, согласовать сметы и графики."),
      ("Коммуникации","План публикаций в соцсетях, встречи с жителями, ответы в комментариях."),
      ("Контроль и KPI","Еженедельный отчёт, дашборд метрик, опрос удовлетворённости.")
    ]
    lines = [f"План действий на месяц — {category} — {municipality_name}",
             f"Период: {today.strftime('%d.%m.%Y')} — {deadline.strftime('%d.%m.%Y')}",
             "", "Цели:", "- Повышение доверия жителей", "- Снижение количества проблемных обращений", "", "Шаги:"]
    for i,(t,d) in enumerate(steps, start=1):
        lines.append(f"{i}. {t}: {d}")
    lines.append("Ключевые KPI: закрытие 80% обращений в срок; рост позитивных упоминаний на 20%; ≥3 встречи с жителями.")
    return "\n".join(lines)

def write_docx(path:str, text:str):
    try:
        import docx
        d = docx.Document()
        for para in text.split("\n"):
            d.add_paragraph(para)
        d.save(path)
        return True
    except Exception:
        return False

def write_pdf(path:str, text:str):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os
        # Locate a Unicode font with Cyrillic support (DejaVu Sans on most Linux images)
        candidates = [
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
        ]
        font_path = next((p for p in candidates if os.path.exists(p)), None)
        if not font_path:
            # Fallback: still attempt with default, but Cyrillic may not render
            font_name = 'Helvetica'
        else:
            font_name = 'DejaVuSans'
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
            except Exception:
                font_name = 'Helvetica'
        c = canvas.Canvas(path, pagesize=A4)
        width, height = A4
        c.setFont(font_name, 11)
        left, top, bottom = 2*cm, height-2*cm, 2*cm
        x, y = left, top
        # Simple word-wrapping
        for raw_line in (text or '').split('\\n'):
            line = str(raw_line).replace('\\t','    ')
            while line:
                # measure characters that fit in line
                n = len(line)
                # naive wrap by characters
                while n>0 and c.stringWidth(line[:n], font_name, 11) > (width - 2*cm - left):
                    n -= 1
                if n == 0:
                    # force break to avoid infinite loop
                    n = 1
                c.drawString(x, y, line[:n])
                line = line[n:].lstrip()
                y -= 14
                if y < bottom:
                    c.showPage()
                    c.setFont(font_name, 11)
                    y = top
        c.save()
        return True
    except Exception as e:
        return False
    except Exception:
        return False

@app.get("/api/appeals/plans")
def list_plans(municipality_id: Optional[int] = None):
    items = DB["plans"]
    if municipality_id:
        items = [p for p in items if p["municipality_id"]==municipality_id]
    return {"items": items[-50:]}

@app.post("/api/appeals/generate-plan/{category}")
async def generate_plan(category:str, payload: dict, request: Request):
    municipality_id = payload.get("municipality_id")
    muni = next((m for m in MUNICIPALITIES if m["id"]==municipality_id), {"name":"Муниципалитет"})
    text = make_plan_text(category, muni["name"])

    plan_id = str(uuid.uuid4())
    docx_path = os.path.join(EXPORT_DIR, f"plan_{plan_id}.docx")
    pdf_path  = os.path.join(EXPORT_DIR, f"plan_{plan_id}.pdf")
    write_docx(docx_path, text)
    write_pdf(pdf_path, text)

    origin = str(request.base_url).rstrip('/')
    item = {
      "id": plan_id,
      "category": category,
      "municipality_id": municipality_id,
      "municipality_name": muni["name"],
      "summary": text.splitlines()[0],
      "created_at": dt.datetime.now().strftime("%Y-%m-%d %H:%M"),
      "docx_url": f"{origin}/api/appeals/file/{os.path.basename(docx_path)}",
      "pdf_url": f"{origin}/api/appeals/file/{os.path.basename(pdf_path)}",
    }
    DB["plans"].append(item)
    return {"ok": True, "item": item}

@app.get("/api/appeals/file/{name}")
@app.get("/appeals/file/{name}")
def get_any_file(name:str):
    path = os.path.join(EXPORT_DIR, name)
    if not os.path.exists(path): raise HTTPException(404, "Файл не найден")
    # simple content-type guess
    mt = "application/octet-stream"
    if name.endswith(".pdf"): mt = "application/pdf"
    if name.endswith(".docx"): mt = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if name.endswith(".xlsx"): mt = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return FileResponse(path, filename=name, media_type=mt)


@app.get('/api/health')
def health():
    return {'ok': True}


# --- Extra fallback CORS middleware (adds headers if something upstream stripped them) ---
class FallbackCORSMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        # Handle preflight
        if request.method == "OPTIONS":
            resp = Response(status_code=204)
        else:
            resp = await call_next(request)

        origin = request.headers.get("origin")
        # If CORSMiddleware already set headers, leave them; otherwise add permissive defaults
        if "access-control-allow-origin" not in (k.lower() for k in resp.headers.keys()):
            # Allow specific origin if provided, else wildcard
            allowed = origin or "*"
            resp.headers["Access-Control-Allow-Origin"] = allowed if allowed != "null" else "*"
            resp.headers["Vary"] = (resp.headers.get("Vary", "") + ", Origin").strip(", ")
            resp.headers["Access-Control-Allow-Methods"] = "GET,POST,PUT,PATCH,DELETE,OPTIONS"
            resp.headers["Access-Control-Allow-Headers"] = request.headers.get("access-control-request-headers", "*") or "*"
            resp.headers["Access-Control-Expose-Headers"] = "Content-Disposition"
        return resp

app.add_middleware(FallbackCORSMiddleware)
def make_plan_text(category:str, municipality_name:str)->str:
    today = dt.date.today()
    deadline = today + dt.timedelta(days=30)
    steps = [
      ("Диагностика проблематики","Собрать первичные данные, верифицировать адресные точки, составить карту очагов."),
      ("Быстрые победы (до 2 недель)","Отработать 2–3 адреса с высокой видимостью; подготовить фото «до/после»."),
      ("Системные меры","Запланировать закупки/МКУ/подрядчики, согласовать сметы и графики."),
      ("Коммуникации","План публикаций в соцсетях, встречи с жителями, ответы в комментариях."),
      ("Контроль и KPI","Еженедельный отчёт, дашборд метрик, опрос удовлетворённости.")
    ]
    lines = [f"План действий на месяц — {category} — {municipality_name}",
             f"Период: {today.strftime('%d.%m.%Y')} — {deadline.strftime('%d.%m.%Y')}",
             "", "Цели:", "- Повышение доверия жителей", "- Снижение количества проблемных обращений", "", "Шаги:"]
    for i,(t,d) in enumerate(steps, start=1):
        lines.append(f"{i}. {t}: {d}")
    lines.append("Ключевые KPI: закрытие 80% обращений в срок; рост позитивных упоминаний на 20%; ≥3 встречи с жителями.")
    return "\n".join(lines)

def write_docx(path:str, text:str):
    try:
        import docx
        d = docx.Document()
        for para in text.split("\n"):
            d.add_paragraph(para)
        d.save(path)
        return True
    except Exception:
        return False

def write_pdf(path:str, text:str):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        c = canvas.Canvas(path, pagesize=A4)
        width, height = A4
        x, y = 2*cm, height-2*cm
        for line in text.split("\n"):
            c.drawString(x, y, line[:120])
            y -= 14
            if y < 2*cm:
                c.showPage()
                y = height-2*cm
        c.save()
        return True
    except Exception:
        return False

@app.get("/api/appeals/plans")
def list_plans(municipality_id: Optional[int] = None):
    items = DB["plans"]
    if municipality_id:
        items = [p for p in items if p["municipality_id"]==municipality_id]
    return {"items": items[-50:]}

@app.post("/api/appeals/generate-plan/{category}")
async def generate_plan(category:str, payload: dict, request: Request):
    municipality_id = payload.get("municipality_id")
    muni = next((m for m in MUNICIPALITIES if m["id"]==municipality_id), {"name":"Муниципалитет"})
    text = make_plan_text(category, muni["name"])

    plan_id = str(uuid.uuid4())
    docx_path = os.path.join(EXPORT_DIR, f"plan_{plan_id}.docx")
    pdf_path  = os.path.join(EXPORT_DIR, f"plan_{plan_id}.pdf")
    write_docx(docx_path, text)
    write_pdf(pdf_path, text)

    origin = str(request.base_url).rstrip('/')
    item = {
      "id": plan_id,
      "category": category,
      "municipality_id": municipality_id,
      "municipality_name": muni["name"],
      "summary": text.splitlines()[0],
      "created_at": dt.datetime.now().strftime("%Y-%m-%d %H:%M"),
      "docx_url": f"{origin}/api/appeals/file/{os.path.basename(docx_path)}",
      "pdf_url": f"{origin}/api/appeals/file/{os.path.basename(pdf_path)}",
    }
    DB["plans"].append(item)
    return {"ok": True, "item": item}

@app.get("/api/appeals/file/{name}")
@app.get("/appeals/file/{name}")
def get_any_file(name:str):
    path = os.path.join(EXPORT_DIR, name)
    if not os.path.exists(path): raise HTTPException(404, "Файл не найден")
    # simple content-type guess
    mt = "application/octet-stream"
    if name.endswith(".pdf"): mt = "application/pdf"
    if name.endswith(".docx"): mt = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if name.endswith(".xlsx"): mt = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return FileResponse(path, filename=name, media_type=mt)


@app.get('/api/health')
def health():
    return {'ok': True}


# --- Extra fallback CORS middleware (adds headers if something upstream stripped them) ---
class FallbackCORSMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        # Handle preflight
        if request.method == "OPTIONS":
            resp = Response(status_code=204)
        else:
            resp = await call_next(request)

        origin = request.headers.get("origin")
        # If CORSMiddleware already set headers, leave them; otherwise add permissive defaults
        if "access-control-allow-origin" not in (k.lower() for k in resp.headers.keys()):
            # Allow specific origin if provided, else wildcard
            allowed = origin or "*"
            resp.headers["Access-Control-Allow-Origin"] = allowed if allowed != "null" else "*"
            resp.headers["Vary"] = (resp.headers.get("Vary", "") + ", Origin").strip(", ")
            resp.headers["Access-Control-Allow-Methods"] = "GET,POST,PUT,PATCH,DELETE,OPTIONS"
            resp.headers["Access-Control-Allow-Headers"] = request.headers.get("access-control-request-headers", "*") or "*"
            resp.headers["Access-Control-Expose-Headers"] = "Content-Disposition"
        return resp

app.add_middleware(FallbackCORSMiddleware)
